from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path("artifacts/uzdream_hosting_beginner_installation_manual.docx")

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
NAVY = "203748"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
PALE_RED = "FDECEC"
RED = "B42318"
PALE_YELLOW = "FFF6D8"
GOLD = "A66B00"
PALE_GREEN = "EAF7EE"
GREEN = "167044"
GRAY = "5A6570"
LIGHT_GRAY = "F4F4F4"
MID_GRAY = "D8DEE5"
WHITE = "FFFFFF"
BLACK = "111111"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Malgun Gothic", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    set_run_font(run, size=8.5, color=GRAY)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), "Malgun Gothic")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color=MID_GRAY, size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True


def add_para(doc: Document, text: str = "", *, bold=False, color=BLACK, size=11, align=None, after=6, before=0, italic=False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)


def add_bullet(doc: Document, text: str, level=0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.keep_together = True
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_together = True
    p.add_run(text)


def add_checkbox(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.first_line_indent = Inches(-0.02)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"□  {text}")
    set_run_font(r, size=10.5)


def add_callout(doc: Document, title: str, body: str, *, fill=PALE_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.3, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc: Document, code: str, *, caption: str | None = None) -> None:
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(caption)
        set_run_font(r, size=9, color=GRAY, bold=True)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(code.strip("\n").splitlines()):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.3, color="24292F")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.6, color=NAVY, bold=True)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            set_run_font(r, size=9.3, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DEEP_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    # Running header and footer.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("UZDREAM.COM  |  서버 호스팅 설치 매뉴얼")
    set_run_font(hr, size=8.5, color=GRAY, bold=True)
    add_bottom_border(hp, color=MID_GRAY, size="6")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("2026-09-02  |  ")
    set_run_font(fr, size=8.5, color=GRAY)
    add_field(fp, "PAGE")
    fr2 = fp.add_run(" / ")
    set_run_font(fr2, size=8.5, color=GRAY)
    add_field(fp, "NUMPAGES")

def add_cover(doc: Document) -> None:
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("초보자용 운영 가이드")
    set_run_font(r, size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("uzdream.com")
    set_run_font(r, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("서버 호스팅 설치 매뉴얼")
    set_run_font(r, size=20, color=DEEP_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("호스팅 개통 · 서버 보안 · 웹 배포 · HTTPS · API 기동 · 점검")
    set_run_font(r, size=11.5, color=GRAY)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(22)
    add_bottom_border(rule, color="AFC7DC", size="14")

    add_para(doc, "작성 기준일  2026년 9월 2일", size=10.5, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "문서 버전  1.0", size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "대상  서버 설치가 처음인 1인 운영자", size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=32)

    add_callout(
        doc,
        "핵심 결론",
        "현재 신청 화면의 메모리 512MB와 Ubuntu 20.04 조합으로는 이 프로젝트의 전체 운영 구성을 설치하지 않습니다. 먼저 Ubuntu 24.04 LTS로 재설치하고 메모리를 최소 2GB, 권장 4GB로 변경합니다. 변경이 불가능하면 이 서버에는 정적 홈페이지와 Nginx만 설치하고 API·DB·Redis·영상 작업은 외부 서비스로 분리합니다.",
        fill=PALE_RED,
        accent=RED,
    )
    add_page_break(doc)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    core = doc.core_properties
    core.title = "uzdream.com 초보자용 서버 호스팅 설치 매뉴얼"
    core.subject = "Ubuntu 호스팅 개통, Nginx, Docker, HTTPS, HanStoneHomepage 배포"
    core.author = "OpenAI Codex"
    core.keywords = "uzdream.com, Ubuntu, Nginx, Docker, Certbot, 서버 설치"
    add_cover(doc)

    add_heading(doc, "0. 이 문서를 사용하는 방법", 1)
    add_para(doc, "명령을 입력하는 위치를 구분하면 실수를 크게 줄일 수 있습니다. 각 코드 상자 위의 ‘내 PC’ 또는 ‘서버’ 표시를 반드시 확인하십시오.")
    add_table(
        doc,
        ["표시", "입력 장소", "예"],
        [
            ["내 PC · PowerShell", "사용자의 Windows 컴퓨터", "웹 빌드, 압축, 서버로 파일 전송"],
            ["서버 · SSH", "원격 Ubuntu 서버 터미널", "Nginx·Docker 설치, 파일 배치, 서비스 기동"],
            ["호스팅 관리자", "제공된 서비스 관리 화면", "결제, OS 재설치, IP 확인, 재부팅"],
            ["도메인 관리자", "도메인 DNS 설정 화면", "A/CNAME 레코드 등록"],
        ],
        [1800, 3000, 4560],
    )
    add_callout(doc, "입력 규칙", "SERVER_IP, SSH_PORT, GITHUB_ID, IMAGE_DIGEST처럼 대문자로 표시된 부분은 예시가 아닙니다. 발급받은 실제 값으로 바꿔 입력하십시오. 괄호나 꺾쇠괄호를 그대로 입력하지 않습니다.", fill=PALE_YELLOW, accent=GOLD)

    add_heading(doc, "개인정보 및 비밀정보", 2)
    add_bullet(doc, "호스팅 화면에 표시된 담당자 이름·이메일·전화번호는 이 문서에 기록하지 않았습니다.")
    add_bullet(doc, "서버 비밀번호, SSH 개인키, DB 비밀번호, OAuth Secret, 결제 Secret Key는 메신저·문서·Git에 넣지 않습니다.")
    add_bullet(doc, "문서의 예시 값은 실제 운영 값이 아닙니다. 운영 환경 파일은 서버의 /etc/hanstone/production.env에만 저장하고 권한을 600으로 제한합니다.")

    add_heading(doc, "전체 순서 한눈에 보기", 2)
    add_table(
        doc,
        ["단계", "작업", "완료 기준"],
        [
            ["1", "호스팅 사양 변경 및 개통", "Ubuntu 24.04 LTS, RAM 2GB 이상, 공인 IP 발급"],
            ["2", "DNS 연결", "uzdream.com과 www가 서버 IP를 조회"],
            ["3", "SSH 접속 및 보안", "일반 관리자 계정, 업데이트, 방화벽, 스왑 완료"],
            ["4", "Nginx·Docker 설치", "버전 확인과 hello-world 성공"],
            ["5", "정적 홈페이지 배포", "HTTP에서 홈페이지 표시"],
            ["6", "HTTPS 적용", "브라우저 자물쇠, 인증서 자동 갱신 시험 성공"],
            ["7", "API 배포", "운영 외부 서비스 준비 후 live·ready 모두 성공"],
            ["8", "최종 점검 및 기록", "체크리스트 완료, 백업·복구·롤백 경로 확인"],
        ],
        [900, 3200, 5260],
    )

    add_page_break(doc)
    add_heading(doc, "1. 현재 호스팅 신청 내용과 설치 가능 여부", 1)
    add_table(
        doc,
        ["항목", "신청 화면의 값", "판정"],
        [
            ["호스팅 ID", "uzdream", "식별용으로 사용"],
            ["도메인", "uzdream.com", "DNS 연결 대상"],
            ["운영체제", "Ubuntu 20.04 (64bit)", "재설치 권장: 표준 보안 지원 종료"],
            ["메모리", "512MB", "전체 서비스 운영 불가; 증설 필요"],
            ["디스크", "30GB", "웹·API 실행 파일은 가능, 영상 원본 보관에는 부족"],
            ["월 트래픽", "320GB", "초기 운영 가능; 사용량 경보 설정"],
            ["공인 IP", "1개", "결제·개통 후 실제 IP 확인 필요"],
            ["현재 상태", "결제/입금 대기, IP·만기일 미표시", "설치 시작 전 개통 완료 필요"],
        ],
        [1800, 3000, 4560],
    )
    add_callout(doc, "중요: Ubuntu 20.04를 그대로 쓰지 않는 이유", "Ubuntu 20.04 LTS의 표준 보안 유지보수는 2025년 5월 종료되었습니다. 현재 Docker Engine 공식 Ubuntu 설치 지원 목록에도 20.04는 포함되지 않습니다. 새 서버라면 Ubuntu 24.04 LTS를 선택하고, 제공되지 않을 때만 22.04 LTS를 선택하십시오.", fill=PALE_RED, accent=RED)

    add_heading(doc, "메모리 512MB에서 가능한 범위", 2)
    add_table(
        doc,
        ["구성", "512MB", "2GB", "4GB 권장"],
        [
            ["Nginx + 정적 홈페이지", "가능", "가능", "가능"],
            ["Nginx + API 1개(외부 DB·Redis)", "권장하지 않음", "제한적 가능", "안정적인 기본안"],
            ["API + DB + Redis + 여러 워커", "불가", "부족할 수 있음", "부하 시험 후 판단"],
            ["ClamAV + FFmpeg/HLS 변환", "불가", "권장하지 않음", "별도 워커 서버 권장"],
        ],
        [3000, 1600, 1800, 2960],
    )
    add_para(doc, "위 메모리 평가는 프로젝트 구성(NestJS API, PostgreSQL, Redis, 메일·영상 워커, ClamAV, FFmpeg)을 기준으로 한 운영 권고입니다. 호스팅사의 보장 수치가 아니며, 최종 용량은 부하 시험으로 확정합니다.", size=9.2, color=GRAY, italic=True)

    add_heading(doc, "권장 운영 구조", 2)
    add_table(
        doc,
        ["사용자", "공개 서버", "외부 관리형 서비스"],
        [
            ["브라우저\nHTTPS 443", "Nginx\n정적 dist/ 제공\n/api → 127.0.0.1:3000 전달", "PostgreSQL · Redis · 객체 저장소/CDN · SMTP\n(영상 변환·악성코드 검사는 별도 워커 권장)"],
        ],
        [1800, 3600, 3960],
    )
    add_callout(doc, "진행 결정", "RAM 증설과 OS 재설치가 끝나기 전에는 2장까지만 진행하십시오. 512MB를 유지해야 한다면 8장 ‘정적 홈페이지 배포’와 9장 ‘HTTPS’까지만 진행하고 10장 API 배포는 중단합니다.", fill=PALE_YELLOW, accent=GOLD)

    add_page_break(doc)
    add_heading(doc, "2. 호스팅 관리자 화면에서 먼저 할 일", 1)
    add_number(doc, "결제 또는 입금을 완료하고 서비스 상태가 ‘사용 중’으로 바뀔 때까지 기다립니다.")
    add_number(doc, "공인 서버 IP가 표시되면 별도 비밀 메모에 기록합니다. 이 문서에는 직접 적지 않는 것을 권장합니다.")
    add_number(doc, "‘OS 재설치’에서 Ubuntu 24.04 LTS 64bit를 선택합니다. 목록에 없으면 22.04 LTS 64bit를 선택합니다.")
    add_number(doc, "‘서비스 변경’ 또는 ‘용량 추가’에서 RAM을 최소 2GB, 가능하면 4GB로 변경합니다.")
    add_number(doc, "‘호스팅 리포트’에서 SSH 포트, 최초 사용자(root 또는 ubuntu), 초기 비밀번호를 확인합니다.")
    add_number(doc, "서버를 켠 뒤 3~10분 기다리고 SSH 접속을 시작합니다.")
    add_callout(doc, "OS 재설치는 데이터 삭제 작업", "운영 데이터를 올린 뒤 OS 재설치를 누르면 서버 파일이 사라질 수 있습니다. 지금처럼 새 서버일 때 먼저 실행하고, 이후에는 백업과 복구 계획 없이 누르지 마십시오.", fill=PALE_RED, accent=RED)

    add_heading(doc, "개통 완료 체크", 2)
    for item in (
        "서비스 상태가 사용 중이다.",
        "서버 IP가 표시된다.",
        "Ubuntu 24.04 LTS 또는 22.04 LTS가 설치됐다.",
        "메모리가 2GB 이상이다(전체 서비스 기준).",
        "SSH 접속 사용자·포트·초기 비밀번호를 확인했다.",
        "호스팅 관리자 화면 로그인에 2단계 인증을 켰다.",
    ):
        add_checkbox(doc, item)

    add_heading(doc, "3. 도메인 DNS 연결", 1)
    add_para(doc, "도메인 관리자의 DNS 설정 화면에서 아래 레코드를 등록합니다. SERVER_IP는 개통 후 표시된 공인 IP로 바꿉니다. 기존에 같은 이름의 A 또는 CNAME 레코드가 있으면 중복으로 추가하지 말고 기존 값을 수정합니다.")
    add_table(
        doc,
        ["종류", "호스트/이름", "값", "TTL"],
        [
            ["A", "@", "SERVER_IP", "300초(초기 전환)"],
            ["CNAME", "www", "uzdream.com", "300초(초기 전환)"],
        ],
        [1200, 1900, 3900, 2360],
    )
    add_para(doc, "DNS 반영은 즉시 끝날 수도 있지만 수 시간 걸릴 수 있습니다. HTTPS 인증서 발급은 두 주소가 모두 새 서버 IP를 가리킨 뒤 진행합니다.")
    add_code(doc, "Resolve-DnsName uzdream.com\nResolve-DnsName www.uzdream.com", caption="내 PC · PowerShell — DNS 확인")
    add_para(doc, "두 결과의 IP가 SERVER_IP와 같으면 완료입니다. 다르면 DNS 값을 다시 확인하고 기다립니다.")

    add_page_break(doc)
    add_heading(doc, "4. Windows에서 서버에 처음 접속하기", 1)
    add_heading(doc, "4.1 SSH 접속", 2)
    add_para(doc, "Windows 시작 메뉴에서 ‘터미널’ 또는 ‘PowerShell’을 엽니다. SSH 포트가 22이면 첫 번째 명령을, 호스팅사가 다른 포트를 안내했다면 두 번째 명령을 사용합니다.")
    add_code(doc, "ssh root@SERVER_IP\n\n# SSH 포트가 22가 아닐 때만 사용\nssh -p SSH_PORT root@SERVER_IP", caption="내 PC · PowerShell")
    add_bullet(doc, "처음 연결할 때 Are you sure you want to continue connecting?이 나오면 IP를 다시 확인한 뒤 yes를 입력합니다.")
    add_bullet(doc, "비밀번호 입력 중에는 화면에 글자나 별표가 표시되지 않습니다. 정상 동작이므로 입력 후 Enter를 누릅니다.")
    add_bullet(doc, "Permission denied가 나오면 사용자명, 비밀번호, SSH 포트를 호스팅 리포트와 비교합니다.")

    add_heading(doc, "4.2 서버 정보 확인", 2)
    add_code(doc, "cat /etc/os-release\nfree -h\ndf -h /\nip addr", caption="서버 · SSH")
    add_table(
        doc,
        ["명령", "정상 기준"],
        [
            ["cat /etc/os-release", "VERSION_ID가 24.04 또는 22.04"],
            ["free -h", "Mem total이 계약한 메모리와 대체로 일치"],
            ["df -h /", "루트 디스크가 약 30GB이고 여유 공간이 있음"],
            ["ip addr", "공인/사설 네트워크 인터페이스가 표시"],
        ],
        [3100, 6260],
    )

    add_heading(doc, "4.3 일반 관리자 계정 만들기", 2)
    add_para(doc, "평소에는 root로 작업하지 않고 hanstone 계정에 sudo 권한을 부여합니다. adduser가 새 비밀번호를 물으면 길고 고유한 비밀번호를 입력합니다. 이름·전화번호 질문은 비워 두고 Enter를 눌러도 됩니다.")
    add_code(doc, "adduser hanstone\nusermod -aG sudo hanstone", caption="서버 · SSH — 최초 root 접속 상태")
    add_para(doc, "현재 창을 닫지 않은 채 새 PowerShell 창을 열어 아래 접속을 시험합니다.")
    add_code(doc, "ssh hanstone@SERVER_IP\n# 포트가 다르면: ssh -p SSH_PORT hanstone@SERVER_IP\nsudo whoami", caption="내 PC · 새 PowerShell")
    add_para(doc, "마지막 결과가 root이면 sudo 권한이 정상입니다. 새 계정 접속을 확인하기 전에는 기존 root 창을 닫거나 root 로그인을 차단하지 마십시오.")

    add_page_break(doc)
    add_heading(doc, "5. 기본 보안과 서버 초기 설정", 1)
    add_heading(doc, "5.1 업데이트와 시간대", 2)
    add_code(doc, "sudo apt update\nsudo apt full-upgrade -y\nsudo timedatectl set-timezone Asia/Seoul\nsudo timedatectl\nsudo reboot", caption="서버 · SSH — hanstone 계정")
    add_para(doc, "재부팅하면 SSH 연결이 끊깁니다. 1~3분 뒤 다시 접속합니다.")

    add_heading(doc, "5.2 방화벽", 2)
    add_callout(doc, "SSH를 먼저 허용", "방화벽을 켜기 전에 실제 SSH 포트를 반드시 허용하십시오. 잘못된 포트를 허용하고 방화벽을 켜면 원격 접속이 끊길 수 있습니다.", fill=PALE_RED, accent=RED)
    add_code(doc, "# SSH가 22번 포트인 경우\nsudo ufw allow 22/tcp\n\n# 호스팅사가 다른 SSH 포트를 준 경우 위 줄 대신 사용\n# sudo ufw allow SSH_PORT/tcp\n\nsudo ufw allow 80/tcp\nsudo ufw allow 443/tcp\nsudo ufw enable\nsudo ufw status verbose", caption="서버 · SSH")
    add_para(doc, "DB 5432, Redis 6379, API 3000 포트는 인터넷에 열지 않습니다. 운영 Compose는 API를 127.0.0.1:3000에만 연결하고 Nginx가 대신 공개합니다.")
    add_callout(doc, "Docker 포트 주의", "Docker가 공개 주소(예: 0.0.0.0:3000)에 포트를 게시하면 UFW 규칙을 우회할 수 있습니다. 이 프로젝트의 production Compose처럼 127.0.0.1에만 바인딩하고, Compose 파일을 임의로 3000:3000으로 바꾸지 마십시오.", fill=PALE_YELLOW, accent=GOLD)

    add_heading(doc, "5.3 2GB 스왑 추가", 2)
    add_para(doc, "스왑은 순간적인 메모리 부족으로 프로세스가 종료되는 위험을 줄이지만 RAM을 대신하지 않습니다. RAM 증설 후에도 작은 서버에는 2GB 스왑을 두는 편이 안전합니다.")
    add_code(doc, "sudo fallocate -l 2G /swapfile\nsudo chmod 600 /swapfile\nsudo mkswap /swapfile\nsudo swapon /swapfile\necho '/swapfile none swap sw 0 0' | sudo tee -a /etc/fstab\nswapon --show\nfree -h", caption="서버 · SSH")

    add_heading(doc, "5.4 자동 보안 업데이트", 2)
    add_code(doc, "sudo apt install -y unattended-upgrades\nsudo dpkg-reconfigure -plow unattended-upgrades", caption="서버 · SSH")
    add_para(doc, "설정 화면에서 Yes를 선택합니다. 커널 업데이트 뒤에는 재부팅이 필요할 수 있으므로 월 1회 점검 시간을 정합니다.")

    add_page_break(doc)
    add_heading(doc, "6. Nginx와 기본 도구 설치", 1)
    add_code(doc, "sudo apt update\nsudo apt install -y nginx git rsync curl ca-certificates unzip snapd\nsudo systemctl enable --now nginx\nsystemctl is-active nginx", caption="서버 · SSH")
    add_para(doc, "마지막 결과가 active이면 정상입니다. 브라우저에서 http://SERVER_IP를 열었을 때 Nginx 기본 화면이 보이면 80번 포트와 방화벽도 정상입니다.")

    add_heading(doc, "7. Docker Engine 설치", 1)
    add_para(doc, "운영 서버에서는 Docker의 공식 apt 저장소를 사용합니다. 인터넷에서 받은 편의 설치 스크립트는 운영 환경에 권장되지 않습니다.")
    add_code(doc, "sudo apt update\nsudo apt install -y ca-certificates curl\nsudo install -m 0755 -d /etc/apt/keyrings\nsudo curl -fsSL https://download.docker.com/linux/ubuntu/gpg -o /etc/apt/keyrings/docker.asc\nsudo chmod a+r /etc/apt/keyrings/docker.asc\n\nsudo tee /etc/apt/sources.list.d/docker.sources > /dev/null <<EOF\nTypes: deb\nURIs: https://download.docker.com/linux/ubuntu\nSuites: $(. /etc/os-release && echo \"${UBUNTU_CODENAME:-$VERSION_CODENAME}\")\nComponents: stable\nArchitectures: $(dpkg --print-architecture)\nSigned-By: /etc/apt/keyrings/docker.asc\nEOF\n\nsudo apt update\nsudo apt install -y docker-ce docker-ce-cli containerd.io docker-buildx-plugin docker-compose-plugin\nsudo systemctl enable --now docker", caption="서버 · SSH")
    add_code(doc, "sudo docker run --rm hello-world\nsudo docker compose version\nsudo systemctl is-active docker", caption="서버 · SSH — 설치 확인")
    add_para(doc, "hello-world의 성공 안내, Docker Compose 버전, active가 차례로 표시되면 완료입니다. 이 문서는 안전을 위해 docker 명령 앞에 sudo를 계속 사용합니다.")

    add_page_break(doc)
    add_heading(doc, "8. 정적 홈페이지 배포", 1)
    add_callout(doc, "512MB 서버도 여기까지는 가능", "정적 홈페이지는 사용자의 PC에서 빌드하고 서버에는 결과물만 올립니다. 서버에서 npm ci 또는 npm run build:web을 실행하지 않으므로 메모리를 절약할 수 있습니다.", fill=PALE_GREEN, accent=GREEN)

    add_heading(doc, "8.1 내 PC에서 배포 파일 만들기", 2)
    add_para(doc, "프로젝트 폴더 F:\\Home Page에서 PowerShell을 열고 아래 명령을 실행합니다. 모든 검증 명령이 오류 없이 끝나야 합니다.")
    add_code(doc, "Set-Location 'F:\\Home Page'\nnpm ci\n$env:WEB_RELEASE_COMMIT_SHA = (git rev-parse HEAD).Trim()\nnpm run build:web\nnpm run manifest:web-deployment\nnpm run verify:web-artifacts\nCompress-Archive -Path .\\dist\\* -DestinationPath .\\hanstone-web.zip -Force", caption="내 PC · PowerShell")
    add_para(doc, "완료 후 프로젝트 폴더에 hanstone-web.zip이 생깁니다. 빌드 실패 상태에서 이전 zip을 재사용하지 마십시오.")

    add_heading(doc, "8.2 서버로 전송", 2)
    add_code(doc, "scp .\\hanstone-web.zip hanstone@SERVER_IP:/tmp/\n# SSH 포트가 다르면: scp -P SSH_PORT .\\hanstone-web.zip hanstone@SERVER_IP:/tmp/", caption="내 PC · PowerShell")

    add_heading(doc, "8.3 릴리스 디렉터리에 설치", 2)
    add_code(doc, "sudo install -d -m 755 /var/www/hanstone/releases/release-001\nsudo unzip -q /tmp/hanstone-web.zip -d /var/www/hanstone/releases/release-001\nsudo chown -R root:root /var/www/hanstone/releases/release-001\nsudo find /var/www/hanstone/releases/release-001 -type d -exec chmod 755 {} \\;\nsudo find /var/www/hanstone/releases/release-001 -type f -exec chmod 644 {} \\;\nsudo ln -sfn /var/www/hanstone/releases/release-001 /var/www/hanstone/current\nls -la /var/www/hanstone/current/", caption="서버 · SSH")
    add_para(doc, "다음 배포는 release-002, release-003처럼 새 폴더를 만든 뒤 current 링크만 바꿉니다. 이전 릴리스 폴더를 한 개 이상 남기면 빠르게 되돌릴 수 있습니다.")

    add_heading(doc, "8.4 Nginx 사이트 설정", 2)
    add_code(doc, "sudo nano /etc/nginx/sites-available/hanstone", caption="서버 · SSH")
    add_para(doc, "편집기가 열리면 아래 내용을 붙여넣습니다. 저장은 Ctrl+O → Enter, 종료는 Ctrl+X입니다.")
    add_code(doc, "server {\n    listen 80;\n    listen [::]:80;\n    server_name uzdream.com www.uzdream.com;\n\n    root /var/www/hanstone/current;\n    index index.html;\n\n    location /api/ {\n        proxy_pass http://127.0.0.1:3000;\n        proxy_http_version 1.1;\n        proxy_set_header Host $host;\n        proxy_set_header X-Real-IP $remote_addr;\n        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n        proxy_set_header X-Forwarded-Proto $scheme;\n    }\n\n    location / {\n        try_files $uri $uri/ /app.html;\n    }\n}\n", caption="/etc/nginx/sites-available/hanstone 내용")
    add_code(doc, "sudo ln -s /etc/nginx/sites-available/hanstone /etc/nginx/sites-enabled/hanstone\n# 아래 명령에서 'No such file'이 나오면 이미 없는 것이므로 다음 단계로 진행\nsudo unlink /etc/nginx/sites-enabled/default\nsudo nginx -t\nsudo systemctl reload nginx", caption="서버 · SSH")
    add_para(doc, "nginx -t 결과에 syntax is ok와 test is successful이 모두 있어야 합니다. 실패하면 reload하지 말고 표시된 파일과 줄 번호를 수정합니다.")

    add_page_break(doc)
    add_heading(doc, "9. HTTPS 인증서 적용", 1)
    add_para(doc, "DNS가 서버 IP를 가리키고 HTTP 홈페이지가 정상 표시된 뒤 실행합니다. Certbot 공식 권장 방식인 snap 패키지를 사용합니다.")
    add_code(doc, "sudo snap install core\nsudo snap refresh core\nsudo snap install --classic certbot\nsudo ln -sf /snap/bin/certbot /usr/local/bin/certbot\nsudo certbot --nginx -d uzdream.com -d www.uzdream.com", caption="서버 · SSH")
    add_bullet(doc, "이메일 주소는 인증서 만료·보안 알림을 받을 운영자 주소를 입력합니다.")
    add_bullet(doc, "약관 동의 질문에는 내용을 확인한 뒤 동의합니다.")
    add_bullet(doc, "HTTP 요청을 HTTPS로 Redirect하는 선택지가 나오면 리디렉션을 선택합니다.")
    add_code(doc, "sudo certbot renew --dry-run\ncurl -I https://uzdream.com\ncurl -I https://www.uzdream.com", caption="서버 · SSH — 갱신 및 접속 확인")
    add_para(doc, "dry-run이 성공하고 curl 결과가 200 또는 정상적인 301/302이면 완료입니다. 브라우저에서도 https://uzdream.com을 열어 자물쇠 아이콘을 확인합니다.")

    add_heading(doc, "10. API 배포 — RAM 2GB 이상에서만", 1)
    add_callout(doc, "시작 전 외부 준비가 필요", "이 프로젝트의 운영 API는 관리형 PostgreSQL, TLS Redis, SMTP, 비공개 객체 저장소/CDN, 법무 승인 값, OAuth·결제 자격증명을 요구합니다. 예시 값을 그대로 둔 상태에서는 정상 기동하지 않도록 설계돼 있습니다.", fill=PALE_RED, accent=RED)

    add_heading(doc, "10.1 필요한 외부 값", 2)
    add_table(
        doc,
        ["구분", "필수 준비", "주의"],
        [
            ["API 이미지", "registry/repository@sha256:… 불변 digest", "latest 태그 사용 금지"],
            ["PostgreSQL", "TLS 연결 URL, 백업/PITR", "이 서버에 직접 설치하지 않음"],
            ["Redis", "rediss:// TLS URL", "6379 공개 금지"],
            ["메일", "SMTP, SPF·DKIM·DMARC", "실제 발송·반송 시험 필요"],
            ["파일/영상", "비공개 객체 저장소, CDN", "30GB 서버 디스크에 원본 저장 금지"],
            ["서비스 연동", "OAuth·토스페이먼츠 운영 키", "승인 전에는 기능 비활성화"],
            ["법무", "정책 버전·승인 시각·문서 SHA-256", "예시 승인값 사용 금지"],
        ],
        [1900, 4200, 3260],
    )

    add_heading(doc, "10.2 배포 파일을 서버로 전송", 2)
    add_code(doc, "scp -r .\\deploy hanstone@SERVER_IP:/tmp/hanstone-deploy\n# 포트가 다르면 scp -P SSH_PORT ...", caption="내 PC · PowerShell — 프로젝트 루트에서")
    add_code(doc, "sudo install -d -m 755 /opt/hanstone\nsudo cp -a /tmp/hanstone-deploy /opt/hanstone/deploy\nsudo install -d -m 700 /etc/hanstone\nsudo install -m 600 /opt/hanstone/deploy/production.env.example /etc/hanstone/production.env\nsudo nano /etc/hanstone/production.env", caption="서버 · SSH")
    add_para(doc, "production.env의 example.com, 000000…, ‘운영_…’ 같은 예시를 모두 실제 값으로 교체합니다. 파일을 화면 공유하거나 복사해 메신저로 보내지 않습니다.")

    add_heading(doc, "10.3 컨테이너 레지스트리 로그인", 2)
    add_para(doc, "API 이미지가 공개 저장소면 이 단계를 생략할 수 있습니다. 비공개 GitHub Container Registry를 사용할 때만 GitHub 사용자명과 Packages 읽기 권한 토큰이 필요합니다. GitHub 협업자 계정은 필요하지 않지만, 저장소 소유 계정 또는 전용 배포 토큰은 필요합니다.")
    add_code(doc, "sudo docker login ghcr.io -u GITHUB_ID", caption="서버 · SSH — 비공개 GHCR 사용 시")
    add_para(doc, "Password에는 GitHub 로그인 비밀번호가 아니라 최소 권한의 토큰을 입력합니다. 명령행에 토큰을 직접 붙이지 않습니다.")

    add_heading(doc, "10.4 마이그레이션·사전 점검·기동", 2)
    add_code(doc, "cd /opt/hanstone\n\nsudo docker compose --env-file /etc/hanstone/production.env -f deploy/compose.production.yaml pull\n\nsudo docker compose --env-file /etc/hanstone/production.env -f deploy/compose.production.yaml run --rm api npm run db:deploy\n\nsudo docker compose --env-file /etc/hanstone/production.env -f deploy/compose.production.yaml run --rm api node dist/production-preflight.js\n\n# 프리플라이트가 모두 pass일 때만 실행\nsudo docker compose --env-file /etc/hanstone/production.env -f deploy/compose.production.yaml up -d", caption="서버 · SSH")
    add_callout(doc, "멈춰야 하는 조건", "db:deploy 또는 production-preflight가 실패하면 up -d를 실행하지 않습니다. 오류 문구와 발생 시각만 기록하고, production.env 원문이나 Secret은 로그에 복사하지 않습니다.", fill=PALE_RED, accent=RED)

    add_heading(doc, "10.5 API 확인", 2)
    add_code(doc, "cd /opt/hanstone\nsudo docker compose --env-file /etc/hanstone/production.env -f deploy/compose.production.yaml ps\n\nsudo docker compose --env-file /etc/hanstone/production.env -f deploy/compose.production.yaml logs --tail=100 api\n\ncurl -fsS http://127.0.0.1:3000/api/v1/health/live\ncurl -fsS http://127.0.0.1:3000/api/v1/health/ready\ncurl -fsS https://uzdream.com/api/v1/health/live\ncurl -fsS https://uzdream.com/api/v1/health/ready", caption="서버 · SSH")
    add_para(doc, "live는 프로세스가 살아 있음을, ready는 DB·Redis 등 준비 상태까지 포함한 판정입니다. 둘 다 성공해야 트래픽을 받습니다.")

    add_page_break(doc)
    add_heading(doc, "11. 최종 홈페이지 시험", 1)
    add_heading(doc, "11.1 기술 점검", 2)
    for item in (
        "https://uzdream.com과 https://www.uzdream.com이 열린다.",
        "주소창에 인증서 경고가 없고 자물쇠가 표시된다.",
        "HTTP 주소가 HTTPS로 자동 이동한다.",
        "API live와 ready가 모두 성공한다(API 배포 시).",
        "브라우저 개발자 도구 Console에 처리되지 않은 오류가 없다.",
        "Network 탭에서 주요 HTML·CSS·JS·이미지가 200 또는 정상 캐시 응답이다.",
        "서버의 free -h, df -h, docker stats --no-stream 결과에 여유가 있다.",
    ):
        add_checkbox(doc, item)
    add_code(doc, "free -h\ndf -h /\nsudo docker stats --no-stream\nsudo journalctl -u nginx --since '30 minutes ago' --no-pager", caption="서버 · SSH")

    add_heading(doc, "11.2 사용자 기능 점검", 2)
    for item in (
        "홈페이지 첫 화면, 상단 메뉴, 본문 링크가 동작한다.",
        "모바일 390×844와 데스크톱 1440×900에서 가로 스크롤이나 겹침이 없다.",
        "회원가입·이메일 인증·로그인·로그아웃이 동작한다.",
        "공개 강의와 바둑 미션을 조회하고 풀이할 수 있다.",
        "학생 진도 저장과 보호자 리포트가 반영된다.",
        "운영자 전용 화면은 일반 사용자에게 차단된다.",
        "결제는 공급자 테스트 모드로 먼저 성공·취소·실패 흐름을 확인했다.",
        "영상 업로드·재생은 객체 저장소와 CDN 주소로 동작한다.",
    ):
        add_checkbox(doc, item)

    add_heading(doc, "12. 업데이트와 되돌리기", 1)
    add_heading(doc, "12.1 정적 홈페이지 새 버전", 2)
    add_number(doc, "내 PC에서 8.1의 빌드·검증·압축을 다시 실행합니다.")
    add_number(doc, "서버에 release-002처럼 새 폴더를 만들고 압축을 풉니다.")
    add_number(doc, "새 폴더 내용을 확인한 뒤 current 링크를 새 폴더로 바꿉니다.")
    add_number(doc, "nginx -t와 홈페이지 점검을 실행합니다.")
    add_number(doc, "문제가 있으면 current 링크를 직전 release-001로 되돌립니다.")
    add_code(doc, "# 새 버전 적용 예\nsudo ln -sfn /var/www/hanstone/releases/release-002 /var/www/hanstone/current\nsudo nginx -t && sudo systemctl reload nginx\n\n# 문제 발생 시 이전 버전으로 복구\nsudo ln -sfn /var/www/hanstone/releases/release-001 /var/www/hanstone/current\nsudo nginx -t && sudo systemctl reload nginx", caption="서버 · SSH")

    add_heading(doc, "12.2 API 새 버전", 2)
    add_para(doc, "production.env의 API_IMAGE를 승인된 새 repository@sha256 digest로 바꾼 뒤 pull → db:deploy → preflight → up -d 순서를 반복합니다. 문제 시 승인된 직전 digest로 되돌립니다. 적용된 DB 마이그레이션은 직접 역실행하지 않고 호환 마이그레이션으로 복구합니다.")

    add_heading(doc, "13. 자주 발생하는 문제", 1)
    add_table(
        doc,
        ["증상", "먼저 확인", "조치"],
        [
            ["SSH 연결 시간 초과", "IP·포트·서버 전원·방화벽", "호스팅 리포트의 실제 포트 확인, 서버 재부팅 후 재시도"],
            ["Permission denied", "사용자명·비밀번호", "root/ubuntu/hanstone 중 제공된 사용자 확인, 비밀번호 재설정"],
            ["사이트가 Nginx 기본 화면", "enabled 사이트와 root", "default 링크 제거, hanstone 링크와 current 경로 확인"],
            ["nginx -t 실패", "표시된 파일·줄 번호", "오타 수정 후 nginx -t 재실행; 성공 전 reload 금지"],
            ["Certbot 인증 실패", "DNS와 80/443 포트", "두 도메인의 IP 일치, 방화벽, HTTP 접속 확인"],
            ["API live 실패", "컨테이너 상태·api 로그", "이미지와 필수 환경변수 확인"],
            ["API ready만 실패", "DB·Redis·마이그레이션", "TLS URL, 접근 허용 IP, db:deploy 결과 확인"],
            ["서버가 멈춤/느림", "free·docker stats·dmesg", "메모리 증설; 512MB라면 API·워커 중단"],
            ["디스크 부족", "df -h, Docker 이미지", "로그 회전·오래된 이미지 점검; 영상은 외부 저장소 사용"],
        ],
        [2600, 2900, 3860],
    )
    add_code(doc, "sudo systemctl status nginx --no-pager\nsudo journalctl -u nginx -n 100 --no-pager\nsudo docker ps -a\nsudo docker compose --env-file /etc/hanstone/production.env -f /opt/hanstone/deploy/compose.production.yaml logs --tail=100\nfree -h\ndf -h /", caption="서버 · SSH — 기본 진단 명령")

    add_page_break(doc)
    add_heading(doc, "14. 운영자가 매주·매월 확인할 일", 1)
    add_table(
        doc,
        ["주기", "확인 항목", "완료 기준"],
        [
            ["매일", "홈페이지·API 상태, 오류 알림", "live/ready 성공, 사용자 오류 급증 없음"],
            ["매주", "디스크·메모리·트래픽, 인증서 타이머", "디스크 80% 미만, OOM 없음, 갱신 타이머 정상"],
            ["매월", "OS/Docker 보안 업데이트, 재부팅", "보안 업데이트 적용 및 서비스 재점검"],
            ["매월", "DB 백업 복원 표본 시험", "격리 환경에서 복원·조회 성공"],
            ["분기", "전체 복구 훈련과 계정 권한", "RPO/RTO 충족, 불필요 계정·토큰 제거"],
            ["변경 때마다", "배포 기록·digest·롤백 대상", "누가/언제/무엇을 배포했는지 추적 가능"],
        ],
        [1400, 4300, 3660],
    )
    add_callout(doc, "백업의 기준", "백업 파일이 ‘있다’는 것만으로는 충분하지 않습니다. 별도 위치에 보관하고, 실제로 복원해 로그인·핵심 데이터 조회가 되는지 정기적으로 확인해야 합니다.", fill=PALE_YELLOW, accent=GOLD)

    add_heading(doc, "15. 최종 완료 체크리스트", 1)
    for item in (
        "호스팅 결제·개통과 공인 IP 발급이 끝났다.",
        "Ubuntu 24.04 LTS(또는 22.04 LTS)와 RAM 2GB 이상을 확인했다.",
        "uzdream.com과 www DNS가 서버 IP를 가리킨다.",
        "일반 관리자 계정과 sudo 접속을 확인했다.",
        "OS 업데이트, UFW 22/80/443, 2GB 스왑을 설정했다.",
        "Nginx와 Docker Engine 설치 검증을 통과했다.",
        "검증된 dist/만 정적 홈페이지로 배포했다.",
        "HTTPS 인증서 발급과 갱신 dry-run을 통과했다.",
        "운영 외부 서비스와 비밀정보를 모두 준비했다(API 배포 시).",
        "db:deploy와 production-preflight 통과 후 API를 기동했다.",
        "live·ready·사용자 흐름·모바일 화면을 확인했다.",
        "직전 정적 릴리스와 API digest의 롤백 경로를 기록했다.",
        "DB·파일 백업 및 복구 시험 일정을 등록했다.",
    ):
        add_checkbox(doc, item)

    add_heading(doc, "16. 공식 참고 자료", 1)
    refs = [
        ("Ubuntu 릴리스 주기 및 지원 기간", "https://ubuntu.com/about/release-cycle"),
        ("Ubuntu 20.04 지원 상태", "https://ubuntu.com/20-04"),
        ("Docker Engine — Ubuntu 공식 설치", "https://docs.docker.com/engine/install/ubuntu/"),
        ("Docker와 방화벽 주의사항", "https://docs.docker.com/engine/install/ubuntu/#firewall-limitations"),
        ("Ubuntu Server — UFW 방화벽", "https://ubuntu.com/server/docs/security-firewall/"),
        ("Nginx Beginner’s Guide", "https://nginx.org/en/docs/beginners_guide.html"),
        ("Certbot — Nginx on Linux 설치", "https://certbot.eff.org/instructions?ws=nginx&os=snap"),
    ]
    for title, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, title, url)
        r = p.add_run(f" — {url}")
        set_run_font(r, size=9, color=GRAY)
    add_para(doc, "공식 자료 확인일: 2026-09-02. 설치 시점에 지원 버전과 명령이 바뀔 수 있으므로 장기 보관 후 다시 사용할 때는 링크의 최신 내용을 확인하십시오.", size=9.2, color=GRAY, italic=True, after=12)

    add_callout(doc, "문서 적용 범위", "이 매뉴얼은 제공된 호스팅 사양과 F:\\Home Page의 HanStoneHomepage 배포 파일을 기준으로 작성했습니다. 호스팅사의 SSH 사용자·포트·OS 목록, 외부 서비스 계약 조건이 다르면 실제 발급값을 우선합니다.", fill=PALE_BLUE, accent=BLUE)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
