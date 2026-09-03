from __future__ import annotations

from pathlib import Path
from shutil import copy2
from copy import deepcopy

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "경영진_개발현황_보고서_2026-08-29.docx"
OUTPUT = ROOT / "reports" / "경영진_개발현황_보고서_2026-08-30.docx"


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_text_preserving_first_run(paragraph: Paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def replace_exact(document: Document, old: str, new: str) -> None:
    matches = [paragraph for paragraph in iter_paragraphs(document) if paragraph.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for {old!r}, found {len(matches)}")
    replace_text_preserving_first_run(matches[0], new)


def replace_fragments(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in iter_paragraphs(document):
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)


def replace_two_part_callout(document: Document, marker: str, heading: str, body: str) -> None:
    matches = [paragraph for paragraph in iter_paragraphs(document) if marker in paragraph.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one callout for {marker!r}, found {len(matches)}")
    paragraph = matches[0]
    newline_index = next((index for index, run in enumerate(paragraph.runs) if "\n" in run.text), None)
    if newline_index is None or newline_index + 1 >= len(paragraph.runs):
        raise RuntimeError(f"Callout does not have two formatted parts: {marker!r}")
    paragraph.runs[0].text = heading + "\n"
    for run in paragraph.runs[1:newline_index + 1]:
        run.text = ""
    paragraph.runs[newline_index + 1].text = body
    for run in paragraph.runs[newline_index + 2:]:
        run.text = ""


def set_progress_row(document: Document, label: str, completed: str, total: str, rate: str) -> None:
    table = document.tables[1]
    for row in table.rows:
        if row.cells[0].text.strip() == label:
            for cell, value in zip(row.cells[1:], (completed, total, rate), strict=True):
                replace_text_preserving_first_run(cell.paragraphs[0], value)
            return
    raise RuntimeError(f"Progress row not found: {label}")


def main() -> None:
    copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    progress_table = document.tables[1]
    completed_rate_format = deepcopy(progress_table.rows[1].cells[3].paragraphs[0].runs[0]._r.rPr)

    replace_exact(document, "기준일: 2026년 8월 29일", "기준일: 2026년 8월 30일")
    replace_exact(
        document,
        "현재 상태: MVP 핵심 기능 구현 완료 / 운영 전 통합 검증 단계",
        "현재 상태: 로컬 기능·품질 검증 완료 / 외부 운영 인수 대기",
    )
    replace_fragments(document, {
        "종합 진척률 87%": "종합 진척률 98%",
        "345개 중 299개 완료(86.7%, 반올림 87%)": "345개 중 337개 완료(97.7%, 반올림 98%)",
        "남은 13%": "남은 2.3%",
    })
    replace_two_part_callout(
        document,
        "345개 중 337개",
        "종합 진척률 98%",
        "전체 QA 체크리스트 345개 중 337개 완료(97.7%, 반올림 98%)입니다. 제품·로컬 품질 검증은 완료 수준이며 남은 2.3%는 운영 SMTP·HTTPS, 법무 승인, 스테이징·후보 배포·closeout·롤백 증빙입니다.",
    )
    replace_exact(
        document,
        "핵심 제품 상태: 학생의 강의 탐색·시청·바둑미션 풀이, 구독 결제, 보호자 리포트, 지도자 수업도우미와 운영 CMS까지 MVP 흐름이 연결됐습니다.",
        "핵심 제품 상태: 학생 학습·바둑미션·결제·보호자·기관 수업·운영 CMS 전 흐름이 연결됐고, 연령대별 계정 제한과 분리 동의까지 구현했습니다.",
    )
    replace_exact(
        document,
        "품질 상태: 기본 실행, 내비게이션, 메인, 미션, 폼, 접근성, 성능 검증 항목은 100% 완료했습니다.",
        "품질 상태: 서버 395건, 웹 77건, 브라우저 59건과 프로덕션 빌드·Prisma·산출물 검사가 통과했고 운영 의존성 취약점은 0건입니다.",
    )
    replace_exact(
        document,
        "출시 판단: 현재는 기능 시연과 제한적 파일럿이 가능한 단계이며, 일반 고객 대상 운영 출시는 P0 운영·보안 검증 완료 후 승인하는 것이 타당합니다.",
        "출시 판단: 제품과 로컬 인수 기준은 충족했습니다. 일반 공개는 운영 Secret·도메인·법무 승인·실배포 증빙 등 외부 8개 항목을 완료한 뒤 승인해야 합니다.",
    )

    replace_exact(
        document,
        "계정·세션·인증 복구·보호자 연결·리포트",
        "연령대·계정 제한·분리 동의·보호자 연결·리포트",
    )
    replace_exact(
        document,
        "25~30분 수업도우미·7종 자료·대형 미션",
        "기관 멤버십·담당 반·수업도우미·7종 자료",
    )
    replace_exact(
        document,
        "React·NestJS·PostgreSQL·CI·E2E·파일 검사",
        "React·NestJS·PostgreSQL·TLS 강제·CI·E2E",
    )

    recent_updates = {
        "연속 클릭 잠금과 동일 clientMoveId 재시도로 중복 답안 제출을 차단했습니다.":
            "만 14세 미만 계정 제한과 아동 계정·학습정보·유료 구독의 분리된 보호자 동의·철회 이력을 구현했습니다.",
        "모달·모바일·지도자 수업 화면이 같은 attemptId와 서버 판 상태를 공유합니다.":
            "API 본문·경로 입력 경계, 세션 쿠키, 민감정보 로그 금지와 운영 DB·Redis·저장소·SMTP 전송 암호화를 자동 검증합니다.",
        "상담 API 오류에도 입력값·동의를 보존하고 같은 내용으로 재시도합니다.":
            "360~1920px에서 긴 한국어 제목·사용자명과 히어로 안전 영역을 검증하고 구형 학습 요약 응답도 안전하게 표시합니다.",
        "자산 지연 로딩·CSS 분리·캐시 최적화로 LCP 568ms, INP 104ms, CLS 0을 기록했습니다.":
            "프로덕션 성능 검사에서 LCP 500ms, INP 80ms, CLS 0을 기록했고 웹·API 운영 의존성 감사는 취약점 0건입니다.",
    }
    for old, new in recent_updates.items():
        replace_exact(document, old, new)
    replace_exact(
        document,
        "기존 정적 화면을 유지하며 React + TypeScript로 기능별 전환 중이고, 서버는 NestJS + Prisma + PostgreSQL 구조입니다. 영상·문서는 비공개 저장소와 단기 서명 URL, 파일 시그니처·ClamAV 검사를 사용하며 결제·OAuth 연동은 서버 컴포넌트로 분리했습니다.",
        "기존 정적 진입 화면과 React + TypeScript 기능 화면을 병행하고, 서버는 NestJS + Prisma + PostgreSQL 구조입니다. 비공개 저장소·서명 URL·ClamAV, OAuth·결제 컴포넌트, 영속 워커와 배포 인수·종료·롤백 검증을 자동화했으며 운영 설정은 암호화 전송을 강제합니다.",
    )

    replace_exact(
        document,
        "전체 작업을 100%로 보고 QA 체크리스트의 345개 항목을 동일 가중치로 산정했습니다. 완료 299개, 미완료 46개로 86.7%이며 경영 보고 수치는 87%로 반올림했습니다. 이 수치는 투입 공수나 예산 집행률이 아니라 ‘기능·품질·운영 검증의 완료율’입니다.",
        "전체 작업을 100%로 보고 QA 체크리스트 345개를 동일 가중치로 산정했습니다. 완료 337개, 미완료 8개로 97.7%이며 경영 보고 수치는 98%로 반올림했습니다. 이는 공수·예산 집행률이 아니라 기능·품질·운영 검증의 완료율입니다.",
    )
    rows = {
        "기본 실행·내비게이션·메인": ("16", "16", "100%"),
        "바둑 데모·실제 바둑미션": ("52", "52", "100%"),
        "모달·폼": ("9", "9", "100%"),
        "회원과 권한": ("28", "29", "97%"),
        "게시판 작성 권한": ("19", "19", "100%"),
        "진도와 보상": ("7", "7", "100%"),
        "강의 CMS·구독 시청 권한": ("37", "37", "100%"),
        "교재·주문·결제": ("20", "20", "100%"),
        "반응형 화면": ("12", "12", "100%"),
        "접근성·성능": ("18", "18", "100%"),
        "보안과 개인정보": ("16", "18", "89%"),
        "프로토타입·출시 검증 기록": ("103", "108", "95%"),
        "전체": ("337", "345", "98%"),
    }
    for label, values in rows.items():
        set_progress_row(document, label, *values)
    for row in progress_table.rows[1:]:
        if row.cells[0].text.strip() != "보안과 개인정보":
            rate_run = row.cells[3].paragraphs[0].runs[0]
            rate_run._r.get_or_add_rPr()
            rate_run._r.replace(rate_run._r.rPr, deepcopy(completed_rate_format))
    replace_exact(
        document,
        "상태 해석\n강점: 핵심 학습 UX·바둑미션·접근성·성능은 완료 기준을 충족했습니다. 주의: 진도·보상, 기관 권한, 개인정보/미성년자 정책, 운영 공급자 실연동은 대규모 유료 출시 전에 반드시 완료해야 합니다.",
        "상태 해석\n강점: 제품 기능과 로컬 품질·보안 자동 검증은 완료 수준입니다. 주의: 잔여 8개는 운영 도메인·Secret·실배포 환경, 법무 승인과 실제 후보 증빙 없이는 완료 처리할 수 없습니다.",
    )

    todo_updates = {
        "운영 외부 연동 확정: 네이버·카카오·Google OAuth 앱, 토스 테스트/운영 키, SMTP 도메인, AWS CloudFront·OAC·키 그룹·도메인을 실제 운영 환경에 연결합니다.":
            "메일·HTTPS: SPF·DKIM·DMARC·반송 웹훅과 실제 도메인·DB·Redis·저장소·SMTP의 암호화 전송 배포 artifact를 확보합니다.",
        "권한 경계 검증: 학생·학부모·지도자·기관 관리자·운영자 메뉴와 API, 지도자 인증·기관 멤버십·담당 반 범위를 실제 계정으로 검증합니다.":
            "법무 승인: 개인정보처리방침과 미성년자·보호자 동의 정책의 서명된 최종 승인 기록을 남깁니다.",
        "보안·개인정보 완료: HTTPS, Secure/HttpOnly/SameSite 쿠키, 입력 검증 전수 점검, 비밀번호·토큰 로그 차단, 개인정보 고지와 미성년자 동의 정책을 확정합니다.":
            "스테이징·운영: 제어 부하·워커 soak와 실제 후보 배포 검증 결과를 90일 증빙으로 보관합니다.",
        "출시 증빙 확보: 스테이징 제어 부하, 워커 soak, 실제 후보 배포 검증, release-closeout.json, 이미지 롤백·재검증 훈련을 수행하고 증빙을 보관합니다.":
            "릴리스 종료: 일곱 증빙·closeout·비운영 롤백 재검증 결과를 변경승인 기록에 첨부합니다.",
        "4.2 P1 - 제품 정합성": "4.2 완료된 주요 로컬 보강",
        "다른 기기 진도 동기화, 보상 중복 방지, 다음 강의 개방, 오답 복습, 주간 학습 지표, 학생-보호자 수치 일치를 검증합니다.":
            "진도 동기화·보상 중복 방지·오답 복습·주간 지표와 학생-보호자 수치 일치 검증을 완료했습니다.",
        "유효 구독의 전체 공개 영상 접근, 종료 구독 내역 보존, 과거 결제 스냅샷, 원본 영상 URL 차단을 검증합니다.":
            "구독 접근·결제 스냅샷·원본 URL 차단과 미성년자 유료 구독의 별도 동의 검증을 완료했습니다.",
        "QR 강의 이동·로그인 복귀·사용/만료 안내와 토스 테스트 결제위젯을 실제 공급자 환경에서 확인합니다.":
            "QR·로그인 복귀, 기관 역할 경계, 개인정보 고지와 감사 이력의 자동 테스트를 완료했습니다.",
        "긴 한국어 문구·사용자명, 제목·버튼 겹침, 히어로 주요 인물 크롭을 대표 기기에서 마감합니다.":
            "긴 한국어 문구·사용자명, 제목·버튼 겹침과 히어로 안전 영역을 6개 대표 해상도에서 검증했습니다.",
        "4.3 P2 - 운영 확장": "4.3 운영 인수 실행 순서",
        "초기 공개 강의 PRE-01 콘텐츠와 CMS 입력 데이터를 확정하고 후속 시대 콘텐츠 제작 책임자를 지정합니다.":
            "1) 운영 계정·도메인·Secret·승인자를 등록하고 법무 승인 문서를 후보에 결합합니다.",
        "첨부파일 전 유형의 확장자·MIME·크기·악성 파일 정책을 운영 샘플로 재검증합니다.":
            "2) 같은 후보로 CI·공급망·브라우저·부하·soak·프리플라이트·복구훈련 증빙을 생성합니다.",
        "관리자 작업과 콘텐츠 수정 이력의 조회·보존 정책을 운영 규정에 반영합니다.":
            "3) 배포 검증·closeout·비운영 롤백 훈련을 승인하고 90일 artifact를 변경기록에 연결합니다.",
    }
    for old, new in todo_updates.items():
        replace_exact(document, old, new)

    decision_rows = {
        "운영 계정·예산": ("운영·스테이징 URL, OAuth·토스·SMTP·저장소·레지스트리 Secret과 승인자 지정", "실환경 증빙 7개 생성 불가"),
        "법무·정책": ("미성년자·보호자 동의와 개인정보처리방침 최종 승인", "아동·유료 사용자 출시 제한"),
        "출시 범위": ("실제 후보 배포 창구, 변경승인자와 90일 증빙 보관 위치 확정", "운영 검증·closeout 진행 불가"),
        "콘텐츠 운영": ("직전 정상 불변 이미지와 격리 롤백 훈련 환경 지정", "장애 시 복구 승인 불가"),
    }
    table = document.tables[2]
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label in decision_rows:
            for cell, value in zip(row.cells[1:], decision_rows[label], strict=True):
                replace_text_preserving_first_run(cell.paragraphs[0], value)

    decision_caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text == "표 3. 경영 의사결정이 필요한 항목"
    )
    decision_caption.paragraph_format.page_break_before = True

    replace_exact(
        document,
        "다음 경영 보고에서는 P0 항목의 증빙 링크, 잔여 46개 QA 항목의 감소 수, 실제 운영 공급자별 성공 여부, 스테이징 장애·복구 결과, 출시 후보 커밋과 승인 기록을 제시합니다. 100%는 코드 작성 완료가 아니라 운영 후보 인수와 롤백 검증까지 완료된 상태로 정의합니다.",
        "다음 경영 보고에서는 잔여 8개 항목의 운영 artifact 링크, 법무 승인, 후보 커밋·이미지 digest, 스테이징 soak, 배포 검증·closeout·롤백 훈련 결과를 제시합니다. 100%는 코드 완료가 아니라 실제 운영 후보 인수와 복구 검증까지 승인된 상태입니다.",
    )
    replace_exact(
        document,
        "docs/QA_CHECKLIST.md - 2026-08-29 완료 299개 / 전체 345개 산정 근거",
        "docs/QA_CHECKLIST.md - 2026-08-30 완료 337개 / 전체 345개 산정 근거",
    )

    document.core_properties.title = "바둑타고 한국사 여행 서비스 개발 현황 보고서"
    document.core_properties.subject = "2026-08-30 경영진 개발 현황 및 잔여 운영 인수 과제"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
