from pathlib import Path
from zipfile import ZipFile

from docx import Document


path = Path("artifacts/uzdream_hosting_beginner_installation_manual.docx")
doc = Document(path)
texts = [p.text for p in doc.paragraphs]
texts.extend(
    p.text
    for table in doc.tables
    for row in table.rows
    for cell in row.cells
    for p in cell.paragraphs
)
blob = "\n".join(texts)
headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

with ZipFile(path) as archive:
    document_xml = archive.read("word/document.xml").decode("utf-8")
    styles_xml = archive.read("word/styles.xml").decode("utf-8")

checks = {
    "literal_plus_after_linebreak_absent": "\n+  " not in blob,
    "quoted_eof_absent": "<<'EOF'" not in blob,
    "unquoted_eof_present": "<<EOF" in blob,
    "pii_email_absent": "@naver.com" not in blob,
    "pii_phone_absent": "010-" not in blob,
    "required_sections_present": all(
        section in blob
        for section in (
            "DNS 연결",
            "Docker Engine 설치",
            "HTTPS 인증서 적용",
            "API 배포",
            "최종 완료 체크리스트",
            "공식 참고 자료",
        )
    ),
    "manual_page_breaks_present": document_xml.count('w:type="page"') >= 8,
    "rows_prevent_split_present": "w:cantSplit" in document_xml,
    "korean_font_declared": "Malgun Gothic" in styles_xml,
    "headings_complete": headings[0].startswith("0.") and headings[-1].startswith("16."),
}

print(f"file={path.resolve()}")
print(f"size_bytes={path.stat().st_size}")
print(
    f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} "
    f"headings={len(headings)} sections={len(doc.sections)}"
)
print(f"first_heading={headings[0]}")
print(f"last_heading={headings[-1]}")
print(f"manual_page_breaks={document_xml.count('w:type=\"page\"')}")
print(f"placeholder_count={sum(blob.count(item) for item in ('SERVER_IP', 'SSH_PORT', 'GITHUB_ID'))}")
for name, passed in checks.items():
    print(f"{name}={'PASS' if passed else 'FAIL'}")

if not all(checks.values()):
    raise SystemExit(1)
